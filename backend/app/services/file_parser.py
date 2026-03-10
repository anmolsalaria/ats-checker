"""File parsing service for PDF and DOCX resume files.

Uses pymupdf (fitz) for PDF extraction — produces clean readable text
without PDF metadata artifacts like << /Filter /FlateDecode >>.
"""

import io
import logging
import re

import fitz  # pymupdf
from docx import Document

logger = logging.getLogger(__name__)

# Patterns that indicate raw PDF metadata leaked into text
_PDF_METADATA_RE = re.compile(
    r"<<\s*/[A-Z].*?>>|"           # << /Filter /FlateDecode >>
    r"\b(?:endobj|endstream|xref|startxref|trailer)\b|"
    r"%PDF-\d|"
    r"/(?:Type|Filter|Length|Linearized|FlateDecode)\b",
    re.IGNORECASE | re.DOTALL,
)


class FileParser:
    """Extracts text from uploaded resume files."""

    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    }

    @staticmethod
    def extract_text(file_content: bytes, content_type: str) -> str:
        """Extract text from file based on content type.

        Args:
            file_content: Raw bytes of the uploaded file.
            content_type: MIME type of the file.

        Returns:
            Extracted text from the document.

        Raises:
            ValueError: If the file type is unsupported or parsing fails.
        """
        file_type = FileParser.SUPPORTED_TYPES.get(content_type)

        if file_type == "pdf":
            return FileParser._extract_from_pdf(file_content)
        elif file_type == "docx":
            return FileParser._extract_from_docx(file_content)
        else:
            raise ValueError(
                f"Unsupported file type: {content_type}. "
                f"Supported types: PDF, DOCX"
            )

    @staticmethod
    def _extract_from_pdf(file_content: bytes) -> str:
        """Extract text from a PDF file using pymupdf (fitz).

        pymupdf produces much cleaner text than PyPDF2:
        - Preserves reading order and layout
        - Handles multi-column resumes
        - Does not leak PDF internal metadata into output
        """
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            text_parts: list[str] = []

            for page in doc:
                # Use "text" mode for clean readable output
                page_text = page.get_text("text")
                if page_text:
                    text_parts.append(page_text)

            doc.close()

            text = "\n".join(text_parts).strip()

            # Strip any residual PDF metadata that might have leaked
            text = _PDF_METADATA_RE.sub("", text)
            # Collapse excessive blank lines
            text = re.sub(r"\n{3,}", "\n\n", text)
            text = text.strip()

            if not text:
                raise ValueError(
                    "Could not extract text from PDF. "
                    "The file may be image-based or encrypted."
                )

            return text

        except ValueError:
            raise
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ValueError(f"Failed to parse PDF file: {str(e)}")

    @staticmethod
    def _extract_from_docx(file_content: bytes) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts: list[str] = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())

            text = "\n".join(text_parts).strip()

            if not text:
                raise ValueError(
                    "Could not extract text from DOCX. The file may be empty."
                )

            return text

        except ValueError:
            raise
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")
